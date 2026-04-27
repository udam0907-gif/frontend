from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from app.core.logging import get_logger

logger = get_logger(__name__)

_BIZNUM_PATTERN = re.compile(r"\b(\d{3}[-\s]?\d{2}[-\s]?\d{5})\b")
_AMOUNT_PATTERN = re.compile(
    r"(?:금액|합계|공급가액|청구금액|지급금액|총액|대금|계)\s*[:：]?\s*[₩￦]?\s*([\d,]+)\s*원?",
    re.IGNORECASE,
)
_AMOUNT_FALLBACK = re.compile(r"[₩￦]?\s*([\d]{4,}(?:,\d{3})*)\s*원")
_DATE_PATTERN = re.compile(r"(\d{4})[.\-년\s](\d{1,2})[.\-월\s](\d{1,2})일?")


class DocumentExtractorService:
    """첨부 서류(PDF/DOCX/XLSX/이미지)에서 검증용 핵심 필드 추출."""

    def extract(self, file_path: str, filename: str) -> dict[str, Any]:
        try:
            text = self._read_text(file_path, filename)
        except Exception as e:
            logger.warning("extractor_read_failed", file=file_path, error=str(e))
            return self._empty()

        return {
            "amount": self._amount(text),
            "vendor_registration_number": self._biznum(text),
            "issue_date": self._date(text),
            "vendor_name": self._vendor_name(text),
            "document_title": self._title(text),
            "raw_text_length": len(text),
        }

    def _read_text(self, file_path: str, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return self._pdf(file_path)
        if ext in (".docx", ".doc"):
            return self._docx(file_path)
        if ext in (".xlsx", ".xls"):
            return self._xlsx(file_path)
        if ext in (".jpg", ".jpeg", ".png"):
            return self._image(file_path)
        return ""

    def _pdf(self, path: str) -> str:
        try:
            import fitz
            doc = fitz.open(path)
            text = "\n".join(p.get_text("text") for p in doc)
            doc.close()
            return text
        except Exception as e:
            logger.warning("pdf_read_failed", path=path, error=str(e))
            return ""

    def _docx(self, path: str) -> str:
        try:
            from docx import Document
            doc = Document(path)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            if p.text.strip():
                                lines.append(p.text)
            return "\n".join(lines)
        except Exception as e:
            logger.warning("docx_read_failed", path=path, error=str(e))
            return ""

    def _xlsx(self, path: str) -> str:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        lines.append("\t".join(cells))
            wb.close()
            return "\n".join(lines)
        except Exception as e:
            logger.warning("xlsx_read_failed", path=path, error=str(e))
            return ""

    def _image(self, path: str) -> str:
        try:
            import pytesseract
            from PIL import Image
            return pytesseract.image_to_string(Image.open(path), lang="kor+eng")
        except Exception as e:
            logger.warning("image_ocr_skipped", path=path, error=str(e))
            return ""

    def _amount(self, text: str) -> float | None:
        m = _AMOUNT_PATTERN.search(text) or _AMOUNT_FALLBACK.search(text)
        if m:
            try:
                return float(m.group(1).replace(",", ""))
            except ValueError:
                pass
        return None

    def _biznum(self, text: str) -> str | None:
        m = _BIZNUM_PATTERN.search(text)
        if m:
            digits = re.sub(r"\D", "", m.group(1))
            if len(digits) == 10:
                return f"{digits[:3]}-{digits[3:5]}-{digits[5:]}"
        return None

    def _date(self, text: str) -> str | None:
        m = _DATE_PATTERN.search(text)
        if m:
            return f"{m.group(1)}-{m.group(2).zfill(2)}-{m.group(3).zfill(2)}"
        return None

    def _vendor_name(self, text: str) -> str | None:
        for pat in [
            re.compile(r"(?:공급자|거래처|업체명|상호)\s*[:：]\s*(.{2,30})"),
            re.compile(r"(?:수신|발신)\s*[:：]\s*(.{2,30})"),
        ]:
            m = pat.search(text)
            if m:
                return m.group(1).strip().split("\n")[0][:50]
        return None

    def _title(self, text: str) -> str | None:
        for line in text.split("\n"):
            s = line.strip()
            if 2 <= len(s) <= 60:
                return s
        return None

    @staticmethod
    def _empty() -> dict[str, Any]:
        return {
            "amount": None,
            "vendor_registration_number": None,
            "issue_date": None,
            "vendor_name": None,
            "document_title": None,
            "raw_text_length": 0,
        }
