"""
Document Generator — 파일 형식별 렌더러

렌더링 분기:
  .docx  → docxtpl ({{placeholder}} 채우기)
  .xlsx  → openpyxl (셀 {{placeholder}} 채우기)
  .pdf   → 원본 복사 (passthrough_copy)
  .jpg/.jpeg/.png → 원본 복사 (passthrough_copy)

원칙: LLM은 helper_text 필드에만 사용. 양식 구조 절대 변경 금지.
"""

from __future__ import annotations

import json
import shutil
import uuid
import zipfile
from decimal import Decimal
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import yaml
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Inches

from app.config import settings
from app.core.exceptions import (
    DocumentGenerationError,
    TemplateError,
    TemplateStructureViolationError,
)
from app.core.logging import get_logger
from app.schemas.docx_schemas import (
    CATEGORY_TO_CHECKBOX,
    DOCX_COMMON_ALIASES,
    DOCX_FIELD_ALIASES,
    DOCX_SCHEMAS,
)
from app.services.llm_service import LLMService

logger = get_logger(__name__)

PROMPTS_DIR = Path(__file__).parent.parent / "prompts"

_DOCX_EXTS = {".docx"}
_XLSX_EXTS = {".xlsx", ".xls"}
_PASSTHROUGH_EXTS = {".pdf", ".jpg", ".jpeg", ".png"}


class DocumentGenerator:
    def __init__(self, llm_service: LLMService) -> None:
        self._llm = llm_service
        self._output_base = Path(settings.storage_documents_path)
        self._output_base.mkdir(parents=True, exist_ok=True)
        self._prompt_config = self._load_prompt_config()

    def _load_prompt_config(self) -> dict[str, Any]:
        config_path = PROMPTS_DIR / "document_helper.yaml"
        try:
            with open(config_path, encoding="utf-8") as f:
                return yaml.safe_load(f)
        except FileNotFoundError:
            logger.warning("document_helper_prompt_not_found", path=str(config_path))
            return {"system": "You are a helpful assistant.", "version": "0.0.0"}

    async def generate(
        self,
        template_path: str,
        field_map: dict[str, Any],
        user_values: dict[str, Any],
        project_data: dict[str, Any],
        expense_item_id: str,
        template_id: str,
        layout_map: dict[str, Any] | None = None,
        document_type: str | None = None,
        render_profile: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """
        파일 형식을 감지하고 적절한 렌더러로 처리한다.

        우선순위:
          XLSX + layout_map 존재 → layout_map 렌더러 (_generate_xlsx_layout)
          XLSX + layout_map 없음 → field_map 렌더러 (_generate_xlsx, 기존)
          DOCX + document_type in DOCX_SCHEMAS → 스키마 렌더러 (_generate_docx_schema, 신규)
          DOCX + 스키마 없음     → field_map 렌더러 (_generate_docx, 기존)
          PDF/이미지              → passthrough_copy

        render_mode 값:
          "excel_rendered"   - 값 입력 완료
          "mapping_needed"   - 셀 매핑 미설정 → 원본 복사
          "passthrough_copy" - PDF/이미지 원본 복사
          "docx_rendered"    - DOCX 렌더링 성공
        """
        if not Path(template_path).exists():
            raise TemplateError(f"템플릿 파일을 찾을 수 없습니다: {template_path}")

        ext = Path(template_path).suffix.lower()

        if ext in _PASSTHROUGH_EXTS:
            output_path = self._copy_passthrough(template_path, expense_item_id, ext)
            return {
                "output_path": output_path,
                "render_mode": "passthrough_copy",
                "generation_trace": self._basic_trace(template_path, template_id, "passthrough_copy"),
            }

        if ext in _XLSX_EXTS:
            if layout_map:
                return self._generate_xlsx_layout(
                    template_path, layout_map, user_values, project_data, expense_item_id, template_id
                )
            return await self._generate_xlsx(
                template_path, field_map, user_values, project_data, expense_item_id, template_id
            )

        # DOCX — 스키마 있으면 스키마 렌더러 우선, 없으면 field_map 렌더러
        if document_type and document_type in DOCX_SCHEMAS:
            ctx = self._build_schema_context(
                document_type,
                user_values,
                project_data,
                DOCX_SCHEMAS[document_type],
            )
            if document_type in {
                "quote",
                "comparative_quote",
                "transaction_statement",
                "expense_resolution",
                "inspection_confirmation",
            }:
                return self._generate_docx_form(
                    template_path,
                    document_type,
                    ctx,
                    expense_item_id,
                    template_id,
                    render_profile=render_profile,
                )
            return self._generate_docx_schema(
                template_path, document_type, user_values, project_data, expense_item_id, template_id
            )
        return await self._generate_docx(
            template_path, field_map, user_values, project_data, expense_item_id, template_id
        )

    # ─── DOCX 스키마 렌더러 (DOCX_SCHEMAS 기반, 신규) ──────────────────────

    def _generate_docx_schema(
        self,
        template_path: str,
        document_type: str,
        user_values: dict[str, Any],
        project_data: dict[str, Any],
        expense_item_id: str,
        template_id: str,
    ) -> dict[str, Any]:
        """
        DOCX_SCHEMAS[document_type] 기반 렌더러.
        1) context 구성 (별칭 적용)
        2) checkbox 필드 ■/□ 처리
        3) line_items 반복 블록 구성
        4) docxtpl 렌더링
        """
        schema = DOCX_SCHEMAS[document_type]
        ctx = self._build_schema_context(document_type, user_values, project_data, schema)
        output_path = self._render_docx(template_path, ctx, expense_item_id)

        trace = {
            "template_path": template_path,
            "template_id": template_id,
            "document_type": document_type,
            "renderer": "docx_schema",
            "render_mode": "docx_rendered",
            "scalar_keys": list(schema.scalar_fields.keys()),
            "checkbox_keys": list(schema.checkbox_fields.keys()),
            "line_items_count": len(ctx.get("line_items", [])),
        }
        return {"output_path": output_path, "render_mode": "docx_rendered", "generation_trace": trace}

    def _build_schema_context(
        self,
        document_type: str,
        user_values: dict[str, Any],
        project_data: dict[str, Any],
        schema: Any,
    ) -> dict[str, Any]:
        """DOCX 스키마 전용 context 구성."""
        # 1) 기본 병합 (project_data < user_values 순)
        ctx: dict[str, Any] = {**project_data, **user_values}

        # 2) 공통 별칭 적용 (덮어쓰지 않고 setdefault)
        for src, dst in DOCX_COMMON_ALIASES.items():
            if src in ctx and dst not in ctx:
                ctx[dst] = ctx[src]

        # 3) 문서 타입별 별칭 적용
        for src, dst in DOCX_FIELD_ALIASES.get(document_type, {}).items():
            if src in ctx and dst not in ctx:
                ctx[dst] = ctx[src]

        # 3-b) project_period 자동 생성 (period_start/end 있고 project_period 없을 때)
        if "project_period" not in ctx:
            ps = ctx.get("period_start", "")
            pe = ctx.get("period_end", "")
            if ps and pe:
                ctx["project_period"] = f"{ps} ~ {pe}"

        if document_type in {"quote", "comparative_quote", "transaction_statement"}:
            if document_type == "comparative_quote":
                recipient_name = str(
                    ctx.get("recipient_name")
                    or ctx.get("vendor_name")
                    or ctx.get("our_company_name")
                    or ""
                ).strip()
                supplier_name = str(ctx.get("compare_vendor_name") or ctx.get("supplier_name") or "").strip()
                supplier_registration = str(
                    ctx.get("compare_vendor_registration")
                    or ctx.get("supplier_registration_number")
                    or ""
                ).strip()
                supplier_contact = str(ctx.get("compare_vendor_contact") or "").strip()
                supplier_address = str(ctx.get("supplier_address") or "").strip()
                supplier_business_type = str(ctx.get("supplier_business_type") or "").strip()
                supplier_business_item = str(ctx.get("supplier_business_item") or "").strip()
                supplier_representative = str(ctx.get("supplier_representative") or "").strip()
                supplier_phone = str(ctx.get("supplier_phone") or "").strip()
                supplier_fax = str(ctx.get("supplier_fax") or "").strip()
                supplier_email = str(ctx.get("supplier_email") or "").strip()
                manager_name = str(
                    ctx.get("our_company_manager_name")
                    or ctx.get("our_company_representative")
                    or ""
                ).strip()
            else:
                recipient_name = str(ctx.get("vendor_name") or ctx.get("recipient_name") or "").strip()
                manager_name = str(
                    ctx.get("our_company_manager_name")
                    or ctx.get("our_company_representative")
                    or ""
                ).strip()
                supplier_name = str(ctx.get("our_company_name") or ctx.get("supplier_name") or "").strip()
                supplier_registration = str(
                    ctx.get("our_company_registration_number")
                    or ctx.get("supplier_registration_number")
                    or ""
                ).strip()
                supplier_address = str(ctx.get("our_company_address") or ctx.get("supplier_address") or "").strip()
                supplier_business_type = str(
                    ctx.get("our_company_business_type") or ctx.get("supplier_business_type") or ""
                ).strip()
                supplier_business_item = str(
                    ctx.get("our_company_business_item") or ctx.get("supplier_business_item") or ""
                ).strip()
                supplier_representative = str(
                    ctx.get("our_company_representative") or ctx.get("supplier_representative") or ""
                ).strip()
                supplier_phone = str(ctx.get("our_company_phone") or "").strip()
                supplier_fax = str(ctx.get("our_company_fax") or "").strip()
                supplier_email = str(ctx.get("our_company_email") or "").strip()
                supplier_contact = " / ".join(
                    part for part in [manager_name, supplier_phone, supplier_email] if part
                )

            recipient_display = recipient_name
            if recipient_display and not recipient_display.endswith("귀하"):
                recipient_display = f"{recipient_display} 귀하"

            ctx["recipient_name"] = recipient_name
            ctx["recipient_display_name"] = recipient_display
            ctx["supplier_name"] = supplier_name
            ctx["supplier_registration_number"] = supplier_registration
            ctx["supplier_address"] = supplier_address
            ctx["supplier_business_type"] = supplier_business_type
            ctx["supplier_business_item"] = supplier_business_item
            ctx["supplier_representative"] = supplier_representative
            ctx["supplier_phone"] = supplier_phone
            ctx["supplier_fax"] = supplier_fax
            ctx["supplier_email"] = supplier_email
            ctx["supplier_contact"] = supplier_contact
            ctx["issue_date"] = str(ctx.get("expense_date") or ctx.get("issue_date") or "").strip()

        # 4) checkbox 필드 ■/□ 주입
        if schema.checkbox_fields:
            category = str(ctx.get("category_type", ""))
            checked_field = CATEGORY_TO_CHECKBOX.get(category, "")
            for field_key in schema.checkbox_fields:
                ctx[field_key] = "■" if field_key == checked_field else ""

        # 5) line_items 반복 블록 구성
        if schema.line_items:
            ctx["line_items"] = self._build_line_items(ctx, schema.line_items)

        return ctx

    def _build_line_items(self, ctx: dict[str, Any], li_spec: Any) -> list[dict[str, Any]]:
        """
        line_items 배열 구성.
        - ctx["line_items"]에 이미 배열이 있으면 그대로 사용
        - 없으면 ctx의 스칼라 값(item_name, quantity 등)으로 단일 행 구성
        - max_rows 초과 시 잘림 + 경고 로그
        """
        rows: list[dict[str, Any]] = ctx.get("line_items") or []

        if not rows:
            row = {}
            for col_key in li_spec.columns:
                val = ctx.get(col_key)
                if val is not None:
                    row[col_key] = val
            if row:
                rows = [row]

        if len(rows) > li_spec.max_rows:
            logger.warning(
                "docx_line_items_truncated",
                provided=len(rows),
                max_rows=li_spec.max_rows,
            )
            rows = rows[: li_spec.max_rows]

        # docxtpl은 빈 셀도 렌더링하므로 누락 키를 빈 문자열로 채움
        # seq는 스키마 컬럼 여부와 무관하게 항상 주입 (NO. 열용)
        columns = list(li_spec.columns.keys())
        return [
            {"seq": i + 1, **{col: row.get(col, "") for col in columns}}
            for i, row in enumerate(rows)
        ]

    def _format_doc_amount(self, value: Any) -> str:
        if value in (None, ""):
            return ""
        return f"{int(Decimal(str(value))):,}"

    def _set_cell_text(self, cell: Any, value: Any) -> None:
        cell.text = "" if value is None else str(value)

    def _find_first_table_with_text(self, doc: Document, needle: str) -> Any | None:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if needle in cell.text:
                        return table
        return None

    # ─── 프로파일 기반 렌더러 ────────────────────────────────────────────────────

    def _dispatch_by_render_profile(
        self,
        doc: Document,
        ctx: dict[str, Any],
        profile: dict[str, Any],
    ) -> None:
        strategy = profile.get("render_strategy", "marker_table")
        if strategy == "paragraph_fill":
            self._fill_by_paragraph_profile(doc, ctx, profile.get("paragraph_config") or {})
        elif strategy == "standard_table":
            self._fill_by_standard_table_profile(doc, ctx, profile.get("standard_table_config") or {})
        elif strategy == "marker_table":
            self._fill_by_marker_table_profile(doc, ctx, profile.get("marker_table_config") or {})
        # docxtpl: _render_docx가 별도 처리하므로 여기선 noop

    def _fill_by_paragraph_profile(
        self,
        doc: Document,
        ctx: dict[str, Any],
        config: dict[str, Any],
    ) -> None:
        """문단 인덱스 기반 채움 (paragraph_fill 전략)."""
        paragraphs = doc.paragraphs
        para_map: dict[str, int] = config.get("paragraph_map") or {}

        for field_key, para_idx in para_map.items():
            if not isinstance(para_idx, int) or para_idx >= len(paragraphs):
                continue
            value = ctx.get(field_key)
            if value is not None:
                paragraphs[para_idx].text = str(value)

        line_items_start = config.get("line_items_para_start")
        if line_items_start is None:
            return
        columns: list[str] = config.get("line_items_columns") or [
            "item_name", "spec", "unit", "quantity", "unit_price", "amount", "remark"
        ]
        items: list[dict[str, Any]] = ctx.get("line_items") or []
        for row_offset, item in enumerate(items):
            for col_offset, col_key in enumerate(columns):
                target_idx = line_items_start + row_offset * len(columns) + col_offset
                if target_idx >= len(paragraphs):
                    break
                val: Any = item.get(col_key, "")
                if col_key in ("unit_price", "amount"):
                    val = self._format_doc_amount(val)
                paragraphs[target_idx].text = str(val)

    def _fill_by_standard_table_profile(
        self,
        doc: Document,
        ctx: dict[str, Any],
        config: dict[str, Any],
    ) -> None:
        """고정 테이블 좌표 기반 채움 (standard_table 전략)."""
        h_idx = config.get("header_table_idx", 0)
        b_idx = config.get("body_table_idx", 1)
        if len(doc.tables) <= max(h_idx, b_idx):
            return

        header_table = doc.tables[h_idx]
        body_table = doc.tables[b_idx]

        recipient = str(ctx.get("recipient_display_name") or ctx.get("recipient_name") or "")
        issue_date = str(ctx.get("issue_date") or "")
        sender_manager = str(ctx.get("our_company_manager_name") or ctx.get("supplier_representative") or "")
        sender_name = str(ctx.get("supplier_name") or "")
        total_amount = self._format_doc_amount(ctx.get("total_amount") or ctx.get("amount") or 0)
        items: list[dict[str, Any]] = ctx.get("line_items") or []

        def _pos(cfg_key: str, default_row: int, default_col: int) -> tuple[int, int]:
            p = config.get(cfg_key) or {}
            return p.get("row", default_row), p.get("col", default_col)

        rr, rc = _pos("recipient_pos", 0, 1)
        dr, dc = _pos("date_pos", 0, 4)
        smr, smc = _pos("sender_manager_pos", 1, 1)
        snr, snc = _pos("sender_name_pos", 2, 1)

        if rr < len(header_table.rows) and rc < len(header_table.rows[rr].cells):
            self._set_cell_text(header_table.rows[rr].cells[rc], f"신 : {recipient}")
        if dr < len(header_table.rows) and dc < len(header_table.rows[dr].cells):
            self._set_cell_text(header_table.rows[dr].cells[dc], f"발 행 일 자: {issue_date}")
        if smr < len(header_table.rows) and smc < len(header_table.rows[smr].cells):
            self._set_cell_text(header_table.rows[smr].cells[smc], f"조 : {sender_manager}" if sender_manager else "조 :")
        if snr < len(header_table.rows) and snc < len(header_table.rows[snr].cells):
            self._set_cell_text(header_table.rows[snr].cells[snc], f"신 : {sender_name}")

        if len(body_table.rows) > 0 and len(body_table.rows[0].cells) > 1:
            self._set_cell_text(
                body_table.rows[0].cells[1],
                f"{recipient} 요청 견적 건" if recipient else "견적 요청 건",
            )
        if len(body_table.rows) > 1 and len(body_table.rows[1].cells) > 1:
            self._set_cell_text(body_table.rows[1].cells[1], f"일금 {total_amount}원" if total_amount else "일금 0원")

        li_start = config.get("line_items_start_row", 3)
        li_end = config.get("line_items_end_row", 7)
        li_cols: dict[str, int] = config.get("line_items_columns") or {
            "seq": 0, "item_name": 1, "spec": 3, "unit_price": 4, "amount": 5
        }

        for row in body_table.rows[li_start:li_end]:
            for cell in row.cells:
                self._set_cell_text(cell, "")

        for idx, (row, item) in enumerate(zip(body_table.rows[li_start:li_end], items), start=1):
            cells = row.cells
            for col_key, col_idx in li_cols.items():
                if col_idx >= len(cells):
                    continue
                if col_key == "seq":
                    self._set_cell_text(cells[col_idx], str(idx))
                elif col_key in ("unit_price", "amount"):
                    self._set_cell_text(cells[col_idx], self._format_doc_amount(item.get(col_key, "")))
                else:
                    self._set_cell_text(cells[col_idx], str(item.get(col_key, "")))

        subtotal_pos = config.get("subtotal_pos")
        if subtotal_pos:
            sr, sc = subtotal_pos.get("row", 7), subtotal_pos.get("col", 5)
            if sr < len(body_table.rows) and sc < len(body_table.rows[sr].cells):
                self._set_cell_text(body_table.rows[sr].cells[sc], total_amount or "0")

        total_pos = config.get("total_pos")
        if total_pos:
            tr, tc = total_pos.get("row", 9), total_pos.get("col", 5)
            if tr < len(body_table.rows) and tc < len(body_table.rows[tr].cells):
                self._set_cell_text(body_table.rows[tr].cells[tc], f"₩{total_amount}" if total_amount else "₩0")

        contact_pos = config.get("contact_pos")
        if contact_pos:
            cr, cc = contact_pos.get("row", 11), contact_pos.get("col", 1)
            sender_phone = str(ctx.get("supplier_phone") or "")
            if cr < len(body_table.rows) and cc < len(body_table.rows[cr].cells):
                self._set_cell_text(body_table.rows[cr].cells[cc], f"{sender_manager} / {sender_phone}".strip(" /"))

        email_pos = config.get("company_email_pos")
        if email_pos:
            er, ec = email_pos.get("row", 12), email_pos.get("col", 1)
            sender_email = str(ctx.get("supplier_email") or "")
            if er < len(body_table.rows) and ec < len(body_table.rows[er].cells):
                self._set_cell_text(body_table.rows[er].cells[ec], f"{sender_name} / {sender_email}".strip(" /"))

    def _fill_by_marker_table_profile(
        self,
        doc: Document,
        ctx: dict[str, Any],
        config: dict[str, Any],
    ) -> None:
        """마커 텍스트 기반 채움 (marker_table 전략)."""
        issue_date = str(ctx.get("issue_date") or "")
        total_amount = self._format_doc_amount(ctx.get("total_amount") or ctx.get("amount") or 0)
        subtotal = self._format_doc_amount(ctx.get("subtotal") or ctx.get("total_amount") or ctx.get("amount") or 0)
        vat = self._format_doc_amount(ctx.get("vat") or 0)
        items: list[dict[str, Any]] = ctx.get("line_items") or []

        date_marker = config.get("date_marker", "작성일자")
        amount_marker = config.get("amount_marker", "합계금액")
        line_marker = config.get("line_items_marker", "품목")
        subtotal_offset = config.get("subtotal_row_offset", 28)
        vat_offset = config.get("vat_row_offset", 29)
        total_offset = config.get("total_row_offset", 30)
        total_col = config.get("total_col", 5)

        issue_table = self._find_first_table_with_text(doc, date_marker)
        if issue_table is not None and issue_table.rows and len(issue_table.rows[0].cells) > 1:
            self._set_cell_text(issue_table.rows[0].cells[1], issue_date)

        amount_table = self._find_first_table_with_text(doc, amount_marker)
        if amount_table is not None and amount_table.rows and len(amount_table.rows[0].cells) >= 4:
            self._set_cell_text(amount_table.rows[0].cells[1], total_amount)
            self._set_cell_text(amount_table.rows[0].cells[3], f"₩ {total_amount}" if total_amount else "")

        line_table = self._find_first_table_with_text(doc, line_marker)
        if line_table is None:
            return

        for row in line_table.rows[1:subtotal_offset]:
            for idx, cell in enumerate(row.cells[1:], start=1):
                self._set_cell_text(cell, "-" if idx == total_col else "")

        for row, item in zip(line_table.rows[1:subtotal_offset], items):
            self._set_cell_text(row.cells[1], item.get("item_name", ""))
            self._set_cell_text(row.cells[2], item.get("spec", ""))
            self._set_cell_text(row.cells[3], item.get("quantity", ""))
            self._set_cell_text(row.cells[4], self._format_doc_amount(item.get("unit_price", "")))
            if len(row.cells) > total_col:
                self._set_cell_text(row.cells[total_col], self._format_doc_amount(item.get("amount", "")))
            if len(row.cells) > total_col + 1:
                self._set_cell_text(row.cells[total_col + 1], item.get("remark", ""))

        if len(line_table.rows) > total_offset:
            self._set_cell_text(line_table.rows[subtotal_offset].cells[total_col], subtotal)
            self._set_cell_text(line_table.rows[vat_offset].cells[total_col], vat)
            self._set_cell_text(line_table.rows[total_offset].cells[total_col - 1], f"₩ {total_amount}" if total_amount else "")
            self._set_cell_text(line_table.rows[total_offset].cells[total_col], f"₩ {total_amount}" if total_amount else "")

    def _fill_quote_like_table(self, doc: Document, ctx: dict[str, Any], issue_date_required: bool) -> None:
        issue_date = str(ctx.get("issue_date") or "")
        total_amount = self._format_doc_amount(ctx.get("total_amount") or ctx.get("amount") or 0)
        subtotal = self._format_doc_amount(ctx.get("subtotal") or ctx.get("total_amount") or ctx.get("amount") or 0)
        vat = self._format_doc_amount(ctx.get("vat") or 0)
        items = ctx.get("line_items") or []

        if issue_date_required:
            issue_table = self._find_first_table_with_text(doc, "작성일자")
            if issue_table is not None and issue_table.rows and len(issue_table.rows[0].cells) > 1:
                self._set_cell_text(issue_table.rows[0].cells[1], issue_date)

        amount_table = self._find_first_table_with_text(doc, "합계금액")
        if amount_table is not None and amount_table.rows and len(amount_table.rows[0].cells) >= 4:
            self._set_cell_text(amount_table.rows[0].cells[1], total_amount)
            self._set_cell_text(amount_table.rows[0].cells[3], f"₩ {total_amount}" if total_amount else "")

        line_table = self._find_first_table_with_text(doc, "품목")
        if line_table is None:
            return

        for row in line_table.rows[1:28]:
            for idx, cell in enumerate(row.cells[1:], start=1):
                if idx == 5:
                    self._set_cell_text(cell, "-")
                else:
                    self._set_cell_text(cell, "")

        for row, item in zip(line_table.rows[1:28], items):
            self._set_cell_text(row.cells[1], item.get("item_name", ""))
            self._set_cell_text(row.cells[2], item.get("spec", ""))
            self._set_cell_text(row.cells[3], item.get("quantity", ""))
            self._set_cell_text(row.cells[4], self._format_doc_amount(item.get("unit_price", "")))
            self._set_cell_text(row.cells[5], self._format_doc_amount(item.get("amount", "")))
            self._set_cell_text(row.cells[6], item.get("remark", ""))

        if len(line_table.rows) > 30:
            self._set_cell_text(line_table.rows[28].cells[5], subtotal)
            self._set_cell_text(line_table.rows[29].cells[5], vat)
            self._set_cell_text(line_table.rows[30].cells[4], f"₩ {total_amount}" if total_amount else "")
            self._set_cell_text(line_table.rows[30].cells[5], f"₩ {total_amount}" if total_amount else "")

    def _fill_standard_quote_table(self, doc: Document, ctx: dict[str, Any]) -> None:
        recipient_display = str(ctx.get("recipient_display_name") or ctx.get("recipient_name") or "")
        issue_date = str(ctx.get("issue_date") or "")
        sender_name = str(ctx.get("supplier_name") or "")
        sender_manager = str(ctx.get("our_company_manager_name") or ctx.get("supplier_representative") or "")
        sender_phone = str(ctx.get("supplier_phone") or "")
        sender_email = str(ctx.get("supplier_email") or "")
        items = ctx.get("line_items") or []
        total_amount = self._format_doc_amount(ctx.get("total_amount") or ctx.get("amount") or 0)

        if len(doc.tables) >= 1 and len(doc.tables[0].rows) >= 3:
            header_table = doc.tables[0]
            self._set_cell_text(header_table.rows[0].cells[1], f"신 : {recipient_display}")
            self._set_cell_text(header_table.rows[0].cells[4], f"발 행 일 자: {issue_date}")
            self._set_cell_text(header_table.rows[1].cells[1], f"조 : {sender_manager}" if sender_manager else "조 :")
            self._set_cell_text(header_table.rows[2].cells[1], f"신 : {sender_name}")
            self._set_cell_text(header_table.rows[2].cells[2], sender_manager)

        if len(doc.tables) < 2 or len(doc.tables[1].rows) < 13:
            return

        body_table = doc.tables[1]
        self._set_cell_text(
            body_table.rows[0].cells[1],
            f"{recipient_display} 요청 견적 건" if recipient_display else "견적 요청 건",
        )
        self._set_cell_text(body_table.rows[1].cells[1], f"일금 {total_amount}원" if total_amount else "일금 0원")

        for row in body_table.rows[3:7]:
            for cell in row.cells:
                self._set_cell_text(cell, "")

        for idx, (row, item) in enumerate(zip(body_table.rows[3:7], items), start=1):
            self._set_cell_text(row.cells[0], str(idx))
            self._set_cell_text(row.cells[1], item.get("item_name", ""))
            self._set_cell_text(row.cells[2], item.get("item_name", ""))
            self._set_cell_text(row.cells[3], item.get("spec") or f"수량 {item.get('quantity', '')}")
            self._set_cell_text(row.cells[4], self._format_doc_amount(item.get("unit_price", "")))
            self._set_cell_text(row.cells[5], self._format_doc_amount(item.get("amount", "")))

        self._set_cell_text(body_table.rows[7].cells[5], total_amount or "0")
        self._set_cell_text(body_table.rows[9].cells[5], f"₩{total_amount}" if total_amount else "₩0")
        self._set_cell_text(body_table.rows[11].cells[1], f"{sender_manager} / {sender_phone}".strip(" /"))
        self._set_cell_text(body_table.rows[12].cells[1], f"{sender_name} / {sender_email}".strip(" /"))

    def _fill_linear_quote_paragraphs(self, doc: Document, ctx: dict[str, Any]) -> None:
        paragraphs = doc.paragraphs
        if len(paragraphs) < 52:
            return

        recipient_display = str(ctx.get("recipient_display_name") or ctx.get("recipient_name") or "")
        issue_date = str(ctx.get("issue_date") or "")
        supplier_name = str(ctx.get("supplier_name") or "")
        supplier_registration = str(ctx.get("supplier_registration_number") or "")
        supplier_business_type = str(ctx.get("supplier_business_type") or "")
        supplier_representative = str(ctx.get("supplier_representative") or "")
        supplier_address = str(ctx.get("supplier_address") or "")
        supplier_phone = str(ctx.get("supplier_phone") or "")
        supplier_fax = str(ctx.get("supplier_fax") or "")
        supplier_business_item = str(ctx.get("supplier_business_item") or "")
        supplier_email = str(ctx.get("supplier_email") or "")
        manager_name = str(ctx.get("our_company_manager_name") or supplier_representative or "")
        total_amount = self._format_doc_amount(ctx.get("total_amount") or ctx.get("amount") or 0)
        items = ctx.get("line_items") or []
        first_item = items[0] if items else {}

        paragraphs[7].text = f"작 성 일\t: {issue_date}"
        paragraphs[12].text = f"{recipient_display}"
        paragraphs[13].text = supplier_name
        paragraphs[15].text = f"공\t등록번호 : {supplier_registration}"
        paragraphs[16].text = f"급\t업태 : {supplier_business_type}\t성명 : {supplier_representative}\t(인)"
        paragraphs[17].text = f"자\t주소 : {supplier_address}"
        paragraphs[18].text = f"전화 : {supplier_phone}\t팩스 : {supplier_fax}"
        paragraphs[19].text = f"{supplier_business_item} 이메일 : {supplier_email}".strip()

        row_values = [
            str(first_item.get("item_name", "")),
            str(first_item.get("spec", "")),
            "EA",
            str(first_item.get("quantity", "")),
            self._format_doc_amount(first_item.get("unit_price", "")),
            self._format_doc_amount(first_item.get("amount", "")),
            str(first_item.get("remark", "")),
        ]
        for idx, value in enumerate(row_values):
            target = 30 + idx
            if target < len(paragraphs):
                paragraphs[target].text = value

        paragraphs[47].text = f"합 계(VAT 제외)\t₩{total_amount}" if total_amount else "합 계(VAT 제외)\t₩0"
        paragraphs[51].text = f"(1) 담당자 : {manager_name}\t휴대폰. {supplier_phone}\te-mail. {supplier_email}".strip()

    def _fill_expense_resolution_table(self, doc: Document, ctx: dict[str, Any]) -> None:
        if len(doc.tables) < 4:
            return

        info_table = doc.tables[1]
        meta_table = doc.tables[2]
        item_table = doc.tables[3]

        project_name = str(ctx.get("project_name") or "")
        project_number = str(ctx.get("project_number") or ctx.get("project_code") or "")
        project_period = str(ctx.get("project_period") or "")
        execution_date = str(ctx.get("execution_date") or ctx.get("expense_date") or "")
        vendor_name = str(ctx.get("vendor_name") or "")
        delivery_date = str(ctx.get("delivery_date") or "")
        usage_purpose = str(ctx.get("usage_purpose") or "")
        purchase_purpose = str(ctx.get("purchase_purpose") or "")
        items = ctx.get("line_items") or []
        total_amount = self._format_doc_amount(ctx.get("total_amount") or ctx.get("amount") or 0)

        self._set_cell_text(info_table.rows[0].cells[1], project_name)
        self._set_cell_text(info_table.rows[1].cells[1], project_number)
        self._set_cell_text(info_table.rows[1].cells[3], project_period)
        self._set_cell_text(info_table.rows[2].cells[1], execution_date)
        self._set_cell_text(info_table.rows[2].cells[3], vendor_name)
        self._set_cell_text(info_table.rows[3].cells[1], delivery_date)

        category_marks = {
            "materials": ["■ 연구재료비", "□ 인건비", "□ 연구활동비", "□ 간접비", "□ 연구수당"],
            "labor": ["□ 연구재료비", "■ 인건비", "□ 연구활동비", "□ 간접비", "□ 연구수당"],
            "outsourcing": ["□ 연구재료비", "□ 인건비", "■ 연구활동비", "□ 간접비", "□ 연구수당"],
            "meeting": ["□ 연구재료비", "□ 인건비", "■ 연구활동비", "□ 간접비", "□ 연구수당"],
            "test_report": ["□ 연구재료비", "□ 인건비", "■ 연구활동비", "□ 간접비", "□ 연구수당"],
            "other": ["□ 연구재료비", "□ 인건비", "□ 연구활동비", "■ 간접비", "□ 연구수당"],
        }
        for idx, mark in enumerate(category_marks.get(str(ctx.get("category_type") or ""), []), start=1):
            self._set_cell_text(meta_table.rows[0].cells[idx], mark)

        self._set_cell_text(meta_table.rows[1].cells[1], usage_purpose)
        self._set_cell_text(meta_table.rows[2].cells[1], purchase_purpose)
        self._set_cell_text(meta_table.rows[3].cells[1], purchase_purpose)

        for row in item_table.rows[1:6]:
            for cell in row.cells[1:]:
                self._set_cell_text(cell, "")

        for idx, (row, item) in enumerate(zip(item_table.rows[1:6], items), start=1):
            self._set_cell_text(row.cells[0], str(idx))
            self._set_cell_text(row.cells[1], item.get("item_name", ""))
            self._set_cell_text(row.cells[2], item.get("spec", ""))
            self._set_cell_text(row.cells[3], item.get("quantity", ""))
            self._set_cell_text(row.cells[4], self._format_doc_amount(item.get("unit_price", "")))
            self._set_cell_text(row.cells[5], self._format_doc_amount(item.get("amount", "")))
            self._set_cell_text(row.cells[6], item.get("remark", ""))

        self._set_cell_text(item_table.rows[6].cells[5], f"{total_amount}원" if total_amount else "0원")

    def _fill_inspection_confirmation_table(self, doc: Document, ctx: dict[str, Any]) -> None:
        if len(doc.tables) < 2:
            return

        main_table = doc.tables[0]
        signer_table = doc.tables[1]

        title = str(ctx.get("contract_name") or ctx.get("title") or "")
        buyer = str(ctx.get("our_company_name") or ctx.get("buyer_name") or "")
        buyer_rep = str(ctx.get("our_company_representative") or "")
        vendor_name = str(ctx.get("vendor_name") or "")
        total_amount = self._format_doc_amount(ctx.get("purchase_amount") or ctx.get("total_amount") or ctx.get("amount") or 0)
        contract_period = str(ctx.get("contract_period") or ctx.get("project_period") or ctx.get("delivery_date") or "")
        delivery_date = str(ctx.get("delivery_date") or "")
        inspection_date = str(ctx.get("inspection_date") or ctx.get("expense_date") or "")
        quantity = str(ctx.get("quantity") or (ctx.get("line_items") or [{}])[0].get("quantity") or "")
        opinion = str(
            ctx.get("inspection_result")
            or ctx.get("inspection_opinion")
            or f"{title or vendor_name} 관련 납품 및 수량을 확인한 결과 이상 없이 완료되었습니다."
        )

        self._set_cell_text(main_table.rows[1].cells[1], title)
        self._set_cell_text(main_table.rows[2].cells[1], f"{buyer} ({buyer_rep})".strip(" ()"))
        self._set_cell_text(main_table.rows[2].cells[5], vendor_name)
        self._set_cell_text(main_table.rows[3].cells[1], f"{total_amount}원" if total_amount else "")
        self._set_cell_text(main_table.rows[3].cells[5], contract_period)
        self._set_cell_text(main_table.rows[4].cells[1], delivery_date)
        self._set_cell_text(main_table.rows[4].cells[5], inspection_date)
        self._set_cell_text(main_table.rows[7].cells[6], f"완료 ({quantity})" if quantity else "완료")
        self._set_cell_text(main_table.rows[8].cells[1], opinion)
        self._set_cell_text(main_table.rows[9].cells[1], "검수 결과 이상 없음")
        self._set_cell_text(main_table.rows[10].cells[1], opinion)

        manager_name = str(ctx.get("our_company_manager_name") or ctx.get("our_company_representative") or "")
        self._set_cell_text(signer_table.rows[0].cells[2], buyer)
        self._set_cell_text(signer_table.rows[1].cells[2], "담당자")
        self._set_cell_text(signer_table.rows[2].cells[2], manager_name)

    def _insert_image_or_clear_marker(
        self,
        doc: Document,
        marker: str,
        image_path: str | None,
    ) -> None:
        def _replace_in_paragraph(paragraph: Any) -> bool:
            if marker not in paragraph.text:
                return False
            for run in paragraph.runs:
                run.text = ""
            if image_path and Path(image_path).exists():
                paragraph.add_run().add_picture(image_path, width=Inches(2.4))
            return True

        for paragraph in doc.paragraphs:
            if _replace_in_paragraph(paragraph):
                return

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if _replace_in_paragraph(paragraph):
                            return

    def _replace_textbox_and_text_runs(self, docx_path: str, ctx: dict[str, Any], document_type: str) -> None:
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        }
        ET.register_namespace("w", ns["w"])
        ET.register_namespace("wps", ns["wps"])

        with zipfile.ZipFile(docx_path, "r") as zf:
            files = {name: zf.read(name) for name in zf.namelist()}

        root = ET.fromstring(files["word/document.xml"])

        recipient_display = str(ctx.get("recipient_display_name") or ctx.get("recipient_name") or "")
        issue_date = str(ctx.get("issue_date") or "")
        supplier_block = " ".join(
            part for part in [
                f"공급자 등록 번호 {ctx.get('supplier_registration_number', '')}".strip(),
                f"상호 {ctx.get('supplier_name', '')}".strip(),
                f"대표자 {ctx.get('supplier_representative', '')} (인)".strip(),
                f"주소 {ctx.get('supplier_address', '')}".strip(),
                f"업태 {ctx.get('supplier_business_type', '')}".strip(),
                f"종목 {ctx.get('supplier_business_item', '')}".strip(),
                f"담당자 {ctx.get('our_company_manager_name', '')}".strip(),
                f"전화 번호 {ctx.get('supplier_phone', '')}".strip(),
                f"팩스 {ctx.get('supplier_fax', '')}".strip(),
                f"이메일 {ctx.get('supplier_email', '')}".strip(),
            ] if part and not part.endswith("()")
        ).strip()

        for txbx in root.findall(".//wps:txbx/w:txbxContent", ns):
            texts = [node.text or "" for node in txbx.findall(".//w:t", ns)]
            joined = "".join(texts).replace(" ", "")
            replacement = None
            if "귀하" in "".join(texts):
                replacement = recipient_display
            elif "작성일자" in "".join(texts):
                replacement = f"작성일자 {issue_date}".strip()
            elif "공급자등록번호" in joined and "상호" in joined:
                replacement = supplier_block

            if replacement is None:
                continue

            first = True
            for node in txbx.findall(".//w:t", ns):
                if first:
                    node.text = replacement
                    first = False
                else:
                    node.text = ""

        if document_type == "transaction_statement":
            for node in root.findall(".//w:t", ns):
                if (node.text or "").strip() == "귀하":
                    node.text = recipient_display

        files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, data in files.items():
                zf.writestr(name, data)

    def _generate_docx_form(
        self,
        template_path: str,
        document_type: str,
        context: dict[str, Any],
        expense_item_id: str,
        template_id: str,
        render_profile: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        # ── docxtpl 전략: Document 직접 조작 없이 docxtpl 렌더러로 위임 ──────
        if render_profile and render_profile.get("render_strategy") == "docxtpl":
            # subtotal/vat 기본값 보장 (docxtpl 템플릿이 해당 변수를 쓸 경우)
            context.setdefault("subtotal", context.get("total_amount", ""))
            context.setdefault("vat", 0)
            output_path = self._render_docx(template_path, context, expense_item_id)
            trace = {
                "template_path": template_path,
                "template_id": template_id,
                "document_type": document_type,
                "renderer": "profile_docxtpl",
                "render_mode": "docx_rendered",
                "fields_filled": list(context.keys()),
                "line_items_count": len(context.get("line_items", [])),
            }
            return {"output_path": output_path, "render_mode": "docx_rendered", "generation_trace": trace}

        output_filename = f"{expense_item_id}_{uuid.uuid4().hex[:8]}.docx"
        output_path = str(self._output_base / output_filename)
        shutil.copy2(template_path, output_path)

        doc = Document(output_path)

        # 프로파일 기반 디스패치 — quote/comparative_quote (marker/paragraph/table 전략)
        if render_profile and document_type in {"quote", "comparative_quote"}:
            self._dispatch_by_render_profile(doc, context, render_profile)
            doc.save(output_path)
            if render_profile.get("textbox_replacement", True):
                self._replace_textbox_and_text_runs(output_path, context, document_type)
            trace = {
                "template_path": template_path,
                "template_id": template_id,
                "document_type": document_type,
                "renderer": f"profile_{render_profile.get('render_strategy', 'unknown')}",
                "render_mode": "docx_rendered",
                "fields_filled": ["supplier_name", "recipient_name", "issue_date", "line_items", "total_amount"],
                "line_items_count": len(context.get("line_items", [])),
            }
            return {"output_path": output_path, "render_mode": "docx_rendered", "generation_trace": trace}

        # 기존 자동감지 fallback (render_profile 없을 때)
        first_table_text = ""
        if doc.tables and doc.tables[0].rows and doc.tables[0].rows[0].cells:
            first_table_text = " ".join(cell.text for cell in doc.tables[0].rows[0].cells)

        if not doc.tables and document_type in {"quote", "comparative_quote"}:
            self._fill_linear_quote_paragraphs(doc, context)
        elif "수 신 :" in first_table_text and "발 행 일 자:" in first_table_text:
            self._fill_standard_quote_table(doc, context)
        elif document_type == "expense_resolution":
            self._fill_expense_resolution_table(doc, context)
        elif document_type == "inspection_confirmation":
            self._fill_inspection_confirmation_table(doc, context)
            self._insert_image_or_clear_marker(
                doc,
                "(이미지)",
                str(context.get("inspection_image_path") or ""),
            )
        else:
            self._fill_quote_like_table(
                doc,
                context,
                issue_date_required=document_type == "transaction_statement",
            )

        doc.save(output_path)
        self._replace_textbox_and_text_runs(output_path, context, document_type)

        trace = {
            "template_path": template_path,
            "template_id": template_id,
            "document_type": document_type,
            "renderer": "docx_form",
            "render_mode": "docx_rendered",
            "fields_filled": [
                "supplier_name",
                "recipient_name",
                "issue_date",
                "line_items",
                "total_amount",
            ],
            "line_items_count": len(context.get("line_items", [])),
        }
        return {"output_path": output_path, "render_mode": "docx_rendered", "generation_trace": trace}

    # ─── DOCX 기존 field_map 렌더러 ({{placeholder}} 방식, legacy) ──────────

    async def _generate_docx(
        self,
        template_path: str,
        field_map: dict[str, Any],
        user_values: dict[str, Any],
        project_data: dict[str, Any],
        expense_item_id: str,
        template_id: str,
    ) -> dict[str, Any]:
        context, llm_fields, token_usage = await self._resolve_fields(
            field_map, user_values, project_data
        )

        output_path = self._render_docx(template_path, context, expense_item_id)

        trace = {
            "template_path": template_path,
            "template_id": template_id,
            "model_version": self._llm._model,
            "prompt_version": self._prompt_config.get("version", "unknown"),
            "render_mode": "docx_rendered",
            "fields_filled": {k: ("***" if "amount" in k else str(v)) for k, v in context.items()},
            "llm_fields": llm_fields,
            "token_usage": token_usage,
            "validation_passed": True,
        }

        return {"output_path": output_path, "render_mode": "docx_rendered", "generation_trace": trace}

    def _render_docx(self, template_path: str, context: dict[str, Any], expense_item_id: str) -> str:
        try:
            tpl = DocxTemplate(template_path)
            safe_context = self._sanitize_context(context)
            tpl.render(safe_context)
            output_filename = f"{expense_item_id}_{uuid.uuid4().hex[:8]}.docx"
            output_path = str(self._output_base / output_filename)
            tpl.save(output_path)
            return output_path
        except Exception as e:
            raise DocumentGenerationError(f"DOCX 렌더링 실패: {e}") from e

    # ─── XLSX layout_map 렌더러 ────────────────────────────────────────────

    def _generate_xlsx_layout(
        self,
        template_path: str,
        layout_map: dict[str, Any],
        user_values: dict[str, Any],
        project_data: dict[str, Any],
        expense_item_id: str,
        template_id: str,
    ) -> dict[str, Any]:
        """
        layout_map 구조(scalar / checkbox / table)를 우선 사용하는 XLSX 렌더러.
        layout_map 없을 때는 호출되지 않는다.
        """
        try:
            import openpyxl
            import openpyxl.utils as _xl_utils
        except ImportError:
            raise DocumentGenerationError("openpyxl이 설치되지 않았습니다.")

        try:
            wb = openpyxl.load_workbook(template_path)
        except Exception as e:
            raise DocumentGenerationError(f"XLSX 파일 열기 실패: {e}") from e

        ws = wb.active
        ctx: dict[str, Any] = {**project_data, **user_values}

        # 병합셀 master 주소 반환 헬퍼
        def _master(addr: str) -> str:
            for mr in ws.merged_cells.ranges:
                for r in range(mr.min_row, mr.max_row + 1):
                    for c in range(mr.min_col, mr.max_col + 1):
                        if f"{_xl_utils.get_column_letter(c)}{r}" == addr:
                            return f"{_xl_utils.get_column_letter(mr.min_col)}{mr.min_row}"
            return addr

        # master → 첫 번째 쓰기 필드 추적 (충돌 방지)
        master_claimed: dict[str, str] = {}
        written: list[str] = []
        skipped: list[str] = []

        def _write(field_key: str, cell_addr: str, value: Any) -> None:
            master = _master(cell_addr)
            if master in master_claimed:
                logger.warning(
                    "layout_xlsx_merge_conflict",
                    field=field_key,
                    cell=cell_addr,
                    master=master,
                    owner=master_claimed[master],
                )
                skipped.append(f"{field_key}(병합충돌→{master_claimed[master]})")
                return
            try:
                ws[master] = value
                master_claimed[master] = field_key
                written.append(f"{field_key}→{master}")
            except Exception as e:
                logger.warning("layout_xlsx_write_failed", field=field_key, cell=master, error=str(e))
                skipped.append(f"{field_key}(쓰기실패:{e})")

        # ── 1. scalar_fields ──────────────────────────────────────────────
        for key, meta in layout_map.get("scalar_fields", {}).items():
            cell = meta.get("cell") if isinstance(meta, dict) else getattr(meta, "cell", None)
            if not cell:
                continue
            value = ctx.get(key)
            if value is None:
                skipped.append(f"{key}(값없음)")
                continue
            _write(key, cell, value)

        # ── 2. checkbox_fields ────────────────────────────────────────────
        for key, meta in layout_map.get("checkbox_fields", {}).items():
            if isinstance(meta, dict):
                cell             = meta.get("cell")
                value_map        = meta.get("value_map", {})
                full_string_mode = meta.get("full_string_mode", True)
                template_string  = meta.get("template_string")
            else:
                cell             = getattr(meta, "cell", None)
                value_map        = getattr(meta, "value_map", {})
                full_string_mode = getattr(meta, "full_string_mode", True)
                template_string  = getattr(meta, "template_string", None)

            if not cell:
                continue

            raw_value = ctx.get(key)
            if raw_value is None:
                skipped.append(f"{key}(값없음)")
                continue

            # context 값 → 체크 레이블
            check_label = value_map.get(str(raw_value), str(raw_value))

            if full_string_mode and template_string:
                # □ {label} → ■ {label} 치환
                cell_value = template_string.replace(f"□ {check_label}", f"■ {check_label}")
            else:
                cell_value = check_label

            _write(key, cell, cell_value)

        # ── 3. table_fields ───────────────────────────────────────────────
        for table_key, meta in layout_map.get("table_fields", {}).items():
            if isinstance(meta, dict):
                start_row = meta.get("start_row", 1)
                max_rows  = meta.get("max_rows", 1)
                columns   = meta.get("columns", {})
            else:
                start_row = getattr(meta, "start_row", 1)
                max_rows  = getattr(meta, "max_rows", 1)
                columns   = getattr(meta, "columns", {})

            # line_items 배열 우선, 없으면 ctx 스칼라 값으로 단일 행 구성
            rows: list[dict[str, Any]] = ctx.get(table_key) or []
            if not rows:
                row = {col: ctx.get(col) for col in columns if ctx.get(col) is not None}
                if row:
                    rows = [row]

            if len(rows) > max_rows:
                logger.warning(
                    "layout_xlsx_table_truncated",
                    table=table_key,
                    provided=len(rows),
                    max_rows=max_rows,
                )
                rows = rows[:max_rows]
                skipped.append(f"{table_key}(행잘림:{len(ctx.get(table_key, []))}→{max_rows})")

            for i, row in enumerate(rows):
                for col_key, col_letter in columns.items():
                    value = row.get(col_key)
                    if value is None:
                        continue
                    cell_addr = f"{col_letter}{start_row + i}"
                    _write(f"{table_key}[{i}].{col_key}", cell_addr, value)

        # ── 저장 ─────────────────────────────────────────────────────────
        output_filename = f"{expense_item_id}_{uuid.uuid4().hex[:8]}.xlsx"
        output_path = str(self._output_base / output_filename)
        try:
            wb.save(output_path)
        except Exception as e:
            raise DocumentGenerationError(f"XLSX 저장 실패: {e}") from e

        render_mode = "excel_rendered" if written else "mapping_needed"
        trace = self._basic_trace(template_path, template_id, render_mode)
        trace["renderer"] = "layout_map"
        trace["xlsx_cells_written"] = len(written)
        trace["written_fields"] = written
        if skipped:
            trace["skipped_fields"] = skipped

        return {"output_path": output_path, "render_mode": render_mode, "generation_trace": trace}

    # ─── XLSX 기존 field_map 렌더러 (layout_map 없을 때 유지) ──────────────

    async def _generate_xlsx(
        self,
        template_path: str,
        field_map: dict[str, Any],
        user_values: dict[str, Any],
        project_data: dict[str, Any],
        expense_item_id: str,
        template_id: str,
    ) -> dict[str, Any]:
        try:
            import openpyxl
        except ImportError:
            raise DocumentGenerationError("openpyxl이 설치되지 않았습니다.")

        # cell 주소가 있는 필드가 하나라도 있는지 확인
        cell_mapped_fields = {
            k: v for k, v in field_map.items()
            if isinstance(v, dict) and v.get("cell")
        }

        if not cell_mapped_fields:
            # 셀 매핑 미설정 → 원본 복사 후 mapping_needed 반환
            import shutil as _shutil
            output_filename = f"{expense_item_id}_{uuid.uuid4().hex[:8]}.xlsx"
            output_path = str(self._output_base / output_filename)
            _shutil.copy2(template_path, output_path)
            trace = self._basic_trace(template_path, template_id, "mapping_needed")
            trace["xlsx_cells_written"] = 0
            return {"output_path": output_path, "render_mode": "mapping_needed", "generation_trace": trace}

        # 컨텍스트 구성: user_values 우선, 없으면 project_data
        context: dict[str, Any] = {**project_data, **user_values}

        try:
            wb = openpyxl.load_workbook(template_path)
        except Exception as e:
            raise DocumentGenerationError(f"XLSX 파일 열기 실패: {e}") from e

        ws = wb.active
        written_count = 0
        written_fields: list[str] = []
        skipped_fields: list[str] = []

        # 병합셀 master 주소 캐시 빌드
        import openpyxl.utils as _xl_utils

        def _master_cell(addr: str) -> str:
            """비master 병합셀이면 좌상단 master 주소를 반환, 아니면 원래 주소."""
            for mr in ws.merged_cells.ranges:
                for r in range(mr.min_row, mr.max_row + 1):
                    for c in range(mr.min_col, mr.max_col + 1):
                        if f"{_xl_utils.get_column_letter(c)}{r}" == addr:
                            master = f"{_xl_utils.get_column_letter(mr.min_col)}{mr.min_row}"
                            return master
            return addr

        # master 셀 → 첫 번째 쓰기 필드만 허용 (중복 충돌 방지)
        master_claimed: dict[str, str] = {}

        for field_key, meta in cell_mapped_fields.items():
            cell_address = meta["cell"]
            value = context.get(field_key)
            if value is None:
                skipped_fields.append(f"{field_key}(값없음)")
                continue

            master = _master_cell(cell_address)
            if master != cell_address:
                logger.info(
                    "xlsx_merged_cell_redirect",
                    field=field_key,
                    configured=cell_address,
                    master=master,
                )

            # 같은 master에 이미 다른 필드가 쓰여졌으면 skip
            if master in master_claimed:
                logger.warning(
                    "xlsx_merged_cell_conflict",
                    field=field_key,
                    cell=cell_address,
                    master=master,
                    already_written_by=master_claimed[master],
                )
                skipped_fields.append(f"{field_key}(병합충돌→{master_claimed[master]})")
                continue

            try:
                ws[master] = value
                master_claimed[master] = field_key
                written_count += 1
                written_fields.append(
                    f"{field_key}→{master}" if master != cell_address
                    else f"{field_key}→{cell_address}"
                )
            except Exception as e:
                logger.warning(
                    "xlsx_cell_write_failed",
                    field=field_key,
                    cell=master,
                    error=str(e),
                )
                skipped_fields.append(f"{field_key}(쓰기실패:{e})")

        output_filename = f"{expense_item_id}_{uuid.uuid4().hex[:8]}.xlsx"
        output_path = str(self._output_base / output_filename)
        try:
            wb.save(output_path)
        except Exception as e:
            raise DocumentGenerationError(f"XLSX 저장 실패: {e}") from e

        render_mode = "excel_rendered" if written_count > 0 else "mapping_needed"
        trace = self._basic_trace(template_path, template_id, render_mode)
        trace["xlsx_cells_written"] = written_count
        trace["written_fields"] = written_fields
        if skipped_fields:
            trace["skipped_fields"] = skipped_fields

        return {"output_path": output_path, "render_mode": render_mode, "generation_trace": trace}

    # ─── Passthrough (PDF / 이미지) ──────────────────────────────────────────

    def _copy_passthrough(self, template_path: str, expense_item_id: str, ext: str) -> str:
        output_filename = f"{expense_item_id}_{uuid.uuid4().hex[:8]}{ext}"
        output_path = str(self._output_base / output_filename)
        shutil.copy2(template_path, output_path)
        return output_path

    # ─── 필드 해석 (DOCX/XLSX 공통) ────────────────────────────────────────

    async def _resolve_fields(
        self,
        field_map: dict[str, Any],
        user_values: dict[str, Any],
        project_data: dict[str, Any],
    ) -> tuple[dict[str, Any], list[str], dict[str, int]]:
        """field_map에 따라 우선순위 기준으로 값을 결정한다."""
        context: dict[str, Any] = {}
        llm_fields: list[str] = []
        token_usage: dict[str, int] = {}

        for placeholder, meta in field_map.items():
            # source는 _meta 또는 최상위 모두 허용 (이전 호환)
            _meta = meta.get("_meta", {})
            source = _meta.get("source") or meta.get("source", "user_input")
            field_type = meta.get("type", "text")

            if placeholder in user_values and user_values[placeholder] is not None:
                context[placeholder] = user_values[placeholder]
                continue

            if source == "project_data" and placeholder in project_data:
                context[placeholder] = project_data[placeholder]
                continue

            if field_type == "helper_text":
                llm_result = await self._fill_helper_text(
                    placeholder=placeholder,
                    label=meta.get("label", placeholder),
                    project_data=project_data,
                    user_values=user_values,
                )
                context[placeholder] = llm_result["text"]
                llm_fields.append(placeholder)
                for k, v in llm_result.get("token_usage", {}).items():
                    token_usage[k] = token_usage.get(k, 0) + v
                continue

            if meta.get("required", True):
                raise DocumentGenerationError(
                    f"필수 항목 누락: {placeholder} ({meta.get('label', placeholder)})"
                )
            context[placeholder] = ""

        # 사용자 값 중 field_map에 없는 키도 컨텍스트에 추가 (업체 원본 양식 지원)
        for k, v in user_values.items():
            if k not in context and v is not None:
                context[k] = v

        return context, llm_fields, token_usage

    # ─── 유틸 ────────────────────────────────────────────────────────────────

    def _sanitize_context(self, context: dict[str, Any]) -> dict[str, Any]:
        import re
        jinja_pattern = re.compile(r"\{[%#].*?[%#]\}", re.DOTALL)
        sanitized = {}
        for k, v in context.items():
            if isinstance(v, str):
                cleaned = jinja_pattern.sub("", v)
                if cleaned != v:
                    raise TemplateStructureViolationError(
                        f"LLM이 템플릿 구조 변경을 시도했습니다. 필드: {k}"
                    )
                sanitized[k] = cleaned
            else:
                sanitized[k] = v
        return sanitized

    def _basic_trace(self, template_path: str, template_id: str, render_mode: str) -> dict:
        return {
            "template_path": template_path,
            "template_id": template_id,
            "render_mode": render_mode,
            "model_version": self._llm._model,
            "prompt_version": self._prompt_config.get("version", "unknown"),
        }

    async def _fill_helper_text(
        self,
        placeholder: str,
        label: str,
        project_data: dict[str, Any],
        user_values: dict[str, Any],
    ) -> dict[str, Any]:
        system = self._prompt_config.get("system", "당신은 R&D 행정 전문가 도우미입니다.")
        prompt_version = self._prompt_config.get("version", "unknown")

        user_msg = (
            f"다음 R&D 과제 서류의 '{label}' 항목을 작성해주세요.\n"
            f"항목 키: {placeholder}\n"
            f"프로젝트 정보: {json.dumps(project_data, ensure_ascii=False)}\n"
            f"사용자 입력 데이터: {json.dumps(user_values, ensure_ascii=False)}\n\n"
            "응답은 해당 항목에 들어갈 텍스트만 반환하세요. "
            "서식이나 마크다운 없이 순수 텍스트로만 작성하세요. "
            "절대로 템플릿 구조({%...%}, {#...#})를 포함하지 마세요."
        )

        response = await self._llm.complete(
            system_prompt=system,
            user_message=user_msg,
            prompt_version=prompt_version,
        )

        return {"text": response.content.strip(), "token_usage": response.token_usage}
