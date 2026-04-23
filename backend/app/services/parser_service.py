from __future__ import annotations

import io
from pathlib import Path
from typing import Any
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from app.core.exceptions import ParseError
from app.core.logging import get_logger

logger = get_logger(__name__)

_WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}


class ParsedPage:
    def __init__(
        self,
        page_number: int,
        text: str,
        section_title: str | None = None,
    ) -> None:
        self.page_number = page_number
        self.text = text
        self.section_title = section_title


class ParserService:
    """
    Parses PDF and DOCX files into text chunks.
    Falls back to OCR for image-based PDFs via pytesseract.
    """

    def parse_file(self, file_path: str, filename: str) -> list[ParsedPage]:
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            return self._parse_xlsx(file_path)
        elif ext in (".jpg", ".jpeg", ".png"):
            return self._parse_image(file_path)
        else:
            raise ParseError(f"지원하지 않는 파일 형식: {ext}")

    def _parse_pdf(self, file_path: str) -> list[ParsedPage]:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ParseError("PyMuPDF가 설치되지 않았습니다.")

        pages: list[ParsedPage] = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text").strip()
                if not text:
                    # Try OCR for image-based pages
                    text = self._ocr_page(page)
                section_title = self._extract_section_title(text)
                pages.append(ParsedPage(
                    page_number=page_num,
                    text=text,
                    section_title=section_title,
                ))
            doc.close()
        except Exception as e:
            raise ParseError(f"PDF 파싱 실패: {e}") from e

        logger.info("pdf_parsed", file=file_path, pages=len(pages))
        return pages

    def _parse_docx(self, file_path: str) -> list[ParsedPage]:
        try:
            from docx import Document
        except ImportError:
            raise ParseError("python-docx가 설치되지 않았습니다.")

        try:
            doc = Document(file_path)
            parts: list[str] = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    cells = [cell for cell in cells if cell]
                    if cells:
                        parts.append("\t".join(cells))

            extra_parts = self._extract_docx_xml_text(file_path)
            for extra in extra_parts:
                if extra and extra not in parts:
                    parts.append(extra)

            full_text = "\n".join(parts)
            chunks = self._split_into_pages(full_text)
            pages = []
            for i, chunk in enumerate(chunks, start=1):
                section_title = self._extract_section_title(chunk)
                pages.append(ParsedPage(
                    page_number=i,
                    text=chunk,
                    section_title=section_title,
                ))
            return pages
        except Exception as e:
            raise ParseError(f"DOCX 파싱 실패: {e}") from e

    def _extract_docx_xml_text(self, file_path: str) -> list[str]:
        parts: list[str] = []
        try:
            with ZipFile(file_path) as archive:
                rel_targets: dict[str, str] = {}
                rels_path = "word/_rels/document.xml.rels"
                if rels_path in archive.namelist():
                    rels_root = ET.fromstring(archive.read(rels_path))
                    for rel in rels_root.findall("rel:Relationship", _WORD_NS):
                        rel_id = rel.attrib.get("Id")
                        target = rel.attrib.get("Target")
                        if rel_id and target:
                            rel_targets[rel_id] = target

                document_path = "word/document.xml"
                if document_path not in archive.namelist():
                    return parts

                root = ET.fromstring(archive.read(document_path))

                for textbox in root.findall(".//w:txbxContent", _WORD_NS):
                    texts = [node.text.strip() for node in textbox.findall(".//w:t", _WORD_NS) if node.text and node.text.strip()]
                    if texts:
                        parts.append(" ".join(texts))

                    for hyperlink in textbox.findall(".//w:hyperlink", _WORD_NS):
                        rel_id = hyperlink.attrib.get(f"{{{_WORD_NS['r']}}}id")
                        if not rel_id:
                            continue
                        target = rel_targets.get(rel_id, "")
                        if target.startswith("mailto:"):
                            parts.append(target.replace("mailto:", "", 1))
        except Exception as e:
            logger.warning("docx_xml_extract_failed", file=file_path, error=str(e))

        return parts

    def _parse_xlsx(self, file_path: str) -> list[ParsedPage]:
        try:
            import openpyxl
        except ImportError:
            raise ParseError("openpyxl이 설치되지 않았습니다. pip install openpyxl")

        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            pages: list[ParsedPage] = []
            for sheet_idx, sheet in enumerate(wb.worksheets, start=1):
                rows_text: list[str] = []
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(cell).strip() if cell is not None else "" for cell in row]
                    if any(c for c in cells):
                        rows_text.append("\t".join(cells))
                if not rows_text:
                    continue
                text = f"[시트: {sheet.title}]\n" + "\n".join(rows_text)
                pages.append(ParsedPage(
                    page_number=sheet_idx,
                    text=text,
                    section_title=sheet.title,
                ))
            wb.close()
            if not pages:
                raise ParseError("XLSX 파일에서 읽을 수 있는 시트가 없습니다.")
            logger.info("xlsx_parsed", file=file_path, sheets=len(pages))
            return pages
        except ParseError:
            raise
        except Exception as e:
            raise ParseError(f"XLSX 파싱 실패: {e}") from e

    def _parse_image(self, file_path: str) -> list[ParsedPage]:
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            raise ParseError("pytesseract 또는 Pillow가 설치되지 않았습니다.")

        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="kor+eng")
            return [ParsedPage(page_number=1, text=text.strip(), section_title=None)]
        except Exception as e:
            raise ParseError(f"이미지 OCR 실패: {e}") from e

    def _ocr_page(self, page: Any) -> str:
        """OCR a PyMuPDF page object."""
        try:
            import pytesseract
            from PIL import Image

            pix = page.get_pixmap(dpi=200)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            return pytesseract.image_to_string(img, lang="kor+eng").strip()
        except Exception as e:
            logger.warning("ocr_failed", error=str(e))
            return ""

    def _split_into_pages(self, text: str, chars_per_page: int = 2000) -> list[str]:
        """Split long DOCX text into pseudo-pages."""
        words = text.split()
        chunks = []
        current: list[str] = []
        current_len = 0

        for word in words:
            current.append(word)
            current_len += len(word) + 1
            if current_len >= chars_per_page:
                chunks.append(" ".join(current))
                current = []
                current_len = 0

        if current:
            chunks.append(" ".join(current))

        return chunks

    def _extract_section_title(self, text: str) -> str | None:
        """Heuristically extract section title from first non-empty line."""
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        if not lines:
            return None
        first_line = lines[0]
        # Title heuristic: short, may start with number
        if len(first_line) <= 100:
            return first_line
        return None

    def chunk_text(
        self,
        pages: list[ParsedPage],
        chunk_size: int = 800,
        chunk_overlap: int = 100,
    ) -> list[dict[str, Any]]:
        """Split pages into overlapping chunks for embedding."""
        chunks = []
        for page in pages:
            text = page.text
            if not text:
                continue

            # Split by characters with overlap
            start = 0
            chunk_idx = 0
            while start < len(text):
                end = min(start + chunk_size, len(text))
                chunk_text = text[start:end].strip()
                if chunk_text:
                    chunks.append({
                        "page_number": page.page_number,
                        "section_title": page.section_title,
                        "chunk_text": chunk_text,
                        "chunk_index": chunk_idx,
                    })
                    chunk_idx += 1
                if end >= len(text):
                    break
                start = end - chunk_overlap

        logger.info("text_chunked", total_chunks=len(chunks))
        return chunks
