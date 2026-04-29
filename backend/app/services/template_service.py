from __future__ import annotations

import re
import shutil
import uuid
from pathlib import Path
from typing import Any

from app.config import settings
from app.core.exceptions import ParseError, StorageError, TemplateError
from app.core.logging import get_logger
from app.core.security import (
    ALLOWED_TEMPLATE_EXTENSIONS,
    generate_safe_filename,
    validate_file_extension,
)

logger = get_logger(__name__)

PLACEHOLDER_PATTERN = re.compile(r"\{\{([^}]+)\}\}")
XLSX_PLACEHOLDER_PATTERN = re.compile(r"\{\{([^}]+)\}\}|\{([^{}]+)\}")


class TemplateService:
    """
    DOCX: docxtpl {{placeholder}} 방식
    XLSX/XLS: openpyxl 셀 값 {{placeholder}} 방식
    Template structure is NEVER modified.
    """

    def __init__(self) -> None:
        self._base_path = Path(settings.storage_templates_path)
        self._base_path.mkdir(parents=True, exist_ok=True)

    def validate_file(self, filename: str, content: bytes) -> None:
        if not validate_file_extension(filename, ALLOWED_TEMPLATE_EXTENSIONS):
            raise TemplateError(
                f"템플릿 파일은 .docx / .xlsx / .xls / .pdf / .jpg / .jpeg / .png 형식만 허용됩니다. 업로드된 파일: {filename}"
            )
        if len(content) > 20 * 1024 * 1024:
            raise TemplateError("템플릿 파일 크기는 20MB를 초과할 수 없습니다.")

    def save_file(self, original_filename: str, content: bytes) -> tuple[str, str]:
        safe_name = generate_safe_filename(original_filename)
        dest_path = self._base_path / safe_name
        try:
            dest_path.write_bytes(content)
        except OSError as e:
            raise StorageError(f"파일 저장 실패: {e}") from e
        logger.info("template_file_saved", path=str(dest_path), size=len(content))
        return safe_name, str(dest_path)

    def extract_placeholders(self, file_path: str) -> dict[str, Any]:
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            return self._extract_docx_placeholders(file_path)
        elif ext in (".xlsx", ".xls"):
            return self._extract_xlsx_placeholders(file_path)
        elif ext in (".pdf", ".jpg", ".jpeg", ".png"):
            return {}
        else:
            raise TemplateError(f"지원하지 않는 템플릿 확장자: {ext}")

    # ── DOCX ──────────────────────────────────────────────────────────────

    def _extract_docx_placeholders(self, file_path: str) -> dict[str, Any]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
        except Exception as e:
            raise ParseError(f"DOCX 파일 파싱 실패: {e}") from e

        placeholders: set[str] = set()
        for para in doc.paragraphs:
            self._scan_text(para.text, placeholders)
        for table in doc.tables:
            self._scan_table_recursive(table, placeholders)
        return self._build_field_map(placeholders)

    def _scan_table_recursive(self, table: Any, placeholders: set[str]) -> None:
        """중첩 테이블 포함 재귀 순회."""
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    self._scan_text(para.text, placeholders)
                for nested_table in cell.tables:
                    self._scan_table_recursive(nested_table, placeholders)

    def _scan_text(self, text: str, placeholders: set[str]) -> None:
        for match in PLACEHOLDER_PATTERN.finditer(text):
            placeholders.add(match.group(1).strip())

    # ── XLSX ──────────────────────────────────────────────────────────────

    def _extract_xlsx_placeholders(self, file_path: str) -> dict[str, Any]:
        try:
            import openpyxl
        except ImportError:
            raise ParseError("openpyxl이 설치되지 않았습니다.")
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=False)
        except Exception as e:
            raise ParseError(f"XLSX 파일 파싱 실패: {e}") from e

        placeholders: set[str] = set()
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if not isinstance(cell.value, str):
                        continue
                    for match in XLSX_PLACEHOLDER_PATTERN.finditer(cell.value):
                        key = (match.group(1) or match.group(2) or "").strip()
                        if key:
                            placeholders.add(key)
        wb.close()
        logger.info("xlsx_placeholders_extracted", file=file_path, count=len(placeholders))
        return self._build_field_map(placeholders)

    # ── layout_map ────────────────────────────────────────────────────────

    def build_layout_map(self, file_path: str) -> dict[str, Any]:
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            return self._docx_layout_map(file_path)
        elif ext in (".xlsx", ".xls"):
            return self._xlsx_layout_map(file_path)
        return {}

    def _docx_layout_map(self, file_path: str) -> dict[str, Any]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
        except Exception:
            return {}
        table_info = []
        for i, tbl in enumerate(doc.tables):
            depth, cells = self._table_depth(tbl)
            table_info.append({
                "index": i,
                "rows": len(tbl.rows),
                "cols": len(tbl.columns),
                "max_depth": depth,
                "cell_count": cells,
            })
        return {
            "format": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "tables": table_info,
        }

    def _table_depth(self, table: Any, _depth: int = 1) -> tuple[int, int]:
        max_depth, cell_count = _depth, 0
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                for nested in cell.tables:
                    d, c = self._table_depth(nested, _depth + 1)
                    max_depth = max(max_depth, d)
                    cell_count += c
        return max_depth, cell_count

    def _xlsx_layout_map(self, file_path: str) -> dict[str, Any]:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            sheets = [
                {"name": s.title, "max_row": s.max_row, "max_column": s.max_column}
                for s in wb.worksheets
            ]
            wb.close()
            return {"format": "xlsx", "sheets": sheets}
        except Exception:
            return {}

    # ── field_map 빌더 ─────────────────────────────────────────────────────

    def _build_field_map(self, placeholders: set[str]) -> dict[str, Any]:
        return {
            p: {
                "label": self._auto_label(p),
                "type": self._auto_type(p),
                "required": not p.startswith("optional_"),
                "source": self._auto_source(p),
                "description": None,
            }
            for p in sorted(placeholders)
        }

    def _auto_label(self, p: str) -> str:
        m = {
            "project_name": "과제명",
            "project_code": "과제번호",
            "institution": "주관기관",
            "pi_name": "연구책임자",
            "expense_date": "지출일",
            "amount": "금액",
            "vendor_name": "거래처명",
            "vendor_registration": "사업자등록번호",
            "description": "내용 설명",
            "period_start": "협약 시작일",
            "period_end": "협약 종료일",
        }
        return m.get(p, p.replace("_", " ").title())

    def _auto_type(self, p: str) -> str:
        if any(k in p for k in ("date", "period", "start", "end")):
            return "date"
        if any(k in p for k in ("amount", "budget", "cost", "price")):
            return "number"
        if any(k in p for k in ("description", "summary", "narrative")):
            return "helper_text"
        return "text"

    def _auto_source(self, p: str) -> str:
        if p in {"project_name", "project_code", "institution", "pi_name", "period_start", "period_end"}:
            return "project_data"
        if self._auto_type(p) == "helper_text":
            return "llm_generated"
        return "user_input"

    def delete_file(self, file_path: str) -> None:
        try:
            Path(file_path).unlink(missing_ok=True)
        except OSError as e:
            logger.warning("template_file_delete_failed", path=file_path, error=str(e))
