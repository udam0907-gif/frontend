"""
docxtpl 호환 DOCX 템플릿 생성 스크립트

대상:
  1. 지출결의서 → /app/storage/templates/expense_resolution_template.docx
  2. 견적서(펀디 기반) → /app/storage/templates/quote_template.docx

원본 파일은 유지하고 새 경로에 저장.
"""
from __future__ import annotations
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree


# ─── 헬퍼 ───────────────────────────────────────────────────────────────────

def set_cell(cell, text: str) -> None:
    """셀의 첫 번째 단락 텍스트를 교체. 셀 서식(테두리·음영)은 유지."""
    para = cell.paragraphs[0]
    # 기존 run 모두 제거
    for run in para.runs:
        run._r.getparent().remove(run._r)
    # 추가 단락 제거 (첫 번째만 유지)
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)
    # 새 run 추가
    para.add_run(text)


def delete_row(table, row_idx: int) -> None:
    """테이블에서 특정 행 삭제."""
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)


def add_empty_row_after(table, after_idx: int, cols: int) -> None:
    """after_idx 행 바로 다음에 빈 행을 복사해 삽입 (endfor 용)."""
    src_tr = table.rows[after_idx]._tr
    new_tr = deepcopy(src_tr)
    # 새 행의 모든 셀을 비움
    for tc in new_tr.findall(f'.//{qn("w:tc")}'):
        for p in tc.findall(f'.//{qn("w:p")}'):
            for r in p.findall(f'{qn("w:r")}'):
                p.remove(r)
            p.append(etree.SubElement(p, qn('w:r')))
    src_tr.addnext(new_tr)
    return new_tr


# ─── 지출결의서 템플릿 ────────────────────────────────────────────────────────

def make_expense_resolution(src: str, dst: str) -> None:
    doc = Document(src)

    # ── table[1]: 과제정보 ──────────────────────────────────────────────────
    t1 = doc.tables[1]
    set_cell(t1.rows[0].cells[1], "{{project_name}}")
    set_cell(t1.rows[1].cells[1], "{{project_number}}")
    set_cell(t1.rows[1].cells[3], "{{project_period}}")
    set_cell(t1.rows[2].cells[1], "{{execution_date}}")
    set_cell(t1.rows[2].cells[3], "{{vendor_name}}")
    set_cell(t1.rows[3].cells[1], "{{delivery_date}}")

    # ── table[2]: 세목 체크박스 + 용도 ────────────────────────────────────
    t2 = doc.tables[2]
    # 세목 체크: '□ 연구재료비' 형태를 '{{var}} 연구재료비' 로 변경
    set_cell(t2.rows[0].cells[1], "{{budget_item_research_materials}} 연구재료비")
    set_cell(t2.rows[0].cells[2], "{{budget_item_labor}} 인건비")
    set_cell(t2.rows[0].cells[3], "{{budget_item_activity}} 연구활동비")
    set_cell(t2.rows[0].cells[4], "{{budget_item_indirect}} 간접비")
    set_cell(t2.rows[0].cells[5], "{{budget_item_allowance}} 연구수당")
    # 용도 / 구매목적
    set_cell(t2.rows[1].cells[1], "{{usage_purpose}}")
    set_cell(t2.rows[2].cells[1], "{{purchase_purpose}}")

    # ── table[3]: 품목표 → line_items 반복 ────────────────────────────────
    # docxtpl {%tr for %} 규칙: for 태그 행 / 콘텐츠 행 / endfor 행 반드시 분리
    #   row[1]: {%tr for %}   → for 태그만, 렌더 후 삭제
    #   row[2]: {{item.xxx}}  → 실제 반복 콘텐츠 (item 변수 사용)
    #   row[3]: {%tr endfor%} → endfor 태그만, 렌더 후 삭제
    #   row[4]: 합산금액
    t3 = doc.tables[3]
    # rows[4-5] 역순 삭제 (row[1-3] 유지 + row[6]→row[4])
    for i in range(5, 3, -1):
        delete_row(t3, i)

    # row[1] → for 태그 행 (콘텐츠 없음)
    for_row = t3.rows[1]
    set_cell(for_row.cells[0], "{%tr for item in line_items %}")
    for ci in range(1, len(for_row.cells)):
        set_cell(for_row.cells[ci], "")

    # row[2] → 콘텐츠 템플릿 행 (item 변수)
    data_row = t3.rows[2]
    set_cell(data_row.cells[0], "{{item.seq}}")
    set_cell(data_row.cells[1], "{{item.item_name}}")
    set_cell(data_row.cells[2], "{{item.spec}}")
    set_cell(data_row.cells[3], "{{item.quantity}}")
    set_cell(data_row.cells[4], "{{item.unit_price}}")
    set_cell(data_row.cells[5], "{{item.amount}}")
    set_cell(data_row.cells[6], "{{item.remark}}")

    # row[3] → endfor 태그 행
    endfor_row = t3.rows[3]
    set_cell(endfor_row.cells[0], "{%tr endfor %}")
    for ci in range(1, len(endfor_row.cells)):
        set_cell(endfor_row.cells[ci], "")

    # row[4] (원 row[6]) → 합산 금액
    total_row = t3.rows[4]
    set_cell(total_row.cells[5], "{{total_amount}}원")
    set_cell(total_row.cells[6], "")

    doc.save(dst)
    print(f"[지출결의서] 저장 완료: {dst}")


# ─── 견적서 템플릿 ────────────────────────────────────────────────────────────

def make_quote(src: str, dst: str) -> None:
    doc = Document(src)

    # ── 헤더 단락: 수신처 / 발행일자 / 발행자 추가 ────────────────────────
    # para[0] = '견적서' 제목, para[1] = 빈 runs
    # para[1]의 첫 번째 run에 수신처/날짜 삽입
    p1 = doc.paragraphs[1]
    for run in p1.runs:
        run._r.getparent().remove(run._r)
    r = p1.add_run("수신 : {{recipient_name}}   발행일자 : {{issue_date}}")
    r.bold = False

    # para[2] = '아래와 같이 견적서를 드립니다.'
    # 다음 단락(para[3])에 공급자 정보 삽입
    p3 = doc.paragraphs[3]
    for run in p3.runs:
        run._r.getparent().remove(run._r)
    p3.add_run("공급자 : {{supplier_name}}   사업자등록번호 : {{supplier_registration_number}}")

    # ── table[0]: 합계금액 ─────────────────────────────────────────────────
    t0 = doc.tables[0]
    # [0][0] = '합계금액... 일금'
    set_cell(t0.rows[0].cells[0], "합계금액 (공급가액) VAT포함   일금")
    # [0][1] = 한글 금액
    set_cell(t0.rows[0].cells[1], "{{total_amount_korean}}")
    set_cell(t0.rows[0].cells[2], "원정")
    # [0][3] = ₩ 숫자
    set_cell(t0.rows[0].cells[3], "₩ {{total_amount}}")

    # ── table[1]: 품목표 → line_items 반복 ────────────────────────────────
    # docxtpl {%tr for %} 규칙: for 태그 행 / 콘텐츠 행 / endfor 행 반드시 분리
    #   row[1]: {%tr for %}   → for 태그만
    #   row[2]: {{item.xxx}}  → 실제 반복 콘텐츠
    #   row[3]: {%tr endfor%} → endfor 태그만
    #   rows[4-6]: 합산/세액/총합계
    t1 = doc.tables[1]
    n_rows = len(t1.rows)  # 31 (header + 27 items + 3 totals)

    # rows[4-27] 역순 삭제 (row[1-3] + rows[28-30] 유지)
    for i in range(n_rows - 4, 3, -1):   # 27 ~ 4 역순
        delete_row(t1, i)

    # row[1] → for 태그 행 (콘텐츠 없음)
    for_row = t1.rows[1]
    set_cell(for_row.cells[0], "{%tr for item in line_items %}")
    for ci in range(1, len(for_row.cells)):
        set_cell(for_row.cells[ci], "")

    # row[2] → 콘텐츠 템플릿 행 (item 변수)
    # 열 순서: NO. | 품목 | 규격 | 수량 | 단가 | 공급가액 | 비고
    data_row = t1.rows[2]
    set_cell(data_row.cells[0], "{{item.seq}}")
    set_cell(data_row.cells[1], "{{item.item_name}}")
    set_cell(data_row.cells[2], "{{item.spec}}")
    set_cell(data_row.cells[3], "{{item.quantity}}")
    set_cell(data_row.cells[4], "{{item.unit_price}}")
    set_cell(data_row.cells[5], "{{item.amount}}")
    if len(data_row.cells) > 6:
        set_cell(data_row.cells[6], "{{item.remark}}")

    # row[3] → endfor 태그 행
    endfor_row = t1.rows[3]
    set_cell(endfor_row.cells[0], "{%tr endfor %}")
    for ci in range(1, len(endfor_row.cells)):
        set_cell(endfor_row.cells[ci], "")

    # rows[4-6]: 합산/세액/총합계
    total_rows = t1.rows
    set_cell(total_rows[4].cells[5], "{{subtotal}}")
    set_cell(total_rows[5].cells[5], "{{vat}}")
    set_cell(total_rows[6].cells[4], "₩ {{total_amount}}")
    set_cell(total_rows[6].cells[5], "₩ {{total_amount}}")

    doc.save(dst)
    print(f"[견적서] 저장 완료: {dst}")


# ─── 실행 ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    BASE = Path("/app/storage")

    make_expense_resolution(
        src=str(BASE / "templates/adc6bd857ca29c8d44564a40a911c447.docx"),
        dst=str(BASE / "templates/expense_resolution_template.docx"),
    )

    make_quote(
        src=str(BASE / "documents/vendors/339cbf90-e167-43d3-91db-1d11dd5c5f78/quote_template_펀디_견적서_양식.docx"),
        dst=str(BASE / "templates/quote_template.docx"),
    )
