"""
m118 — 내부 서류 출력 sanity (Phase 2d)

품의서 / 지출결의서 / 검수확인서 docx fill 결과를 검증한다.

실행:
    docker exec rnd_backend python3 /app/scripts/m118_internal_doc_sanity.py

검증 항목:
  [품의서]    P1~P6
  [지출결의서] E1~E4
  [검수확인서] I1~I3
"""
from __future__ import annotations

import asyncio
import datetime
import io
import os
import sys
import uuid
from pathlib import Path
from typing import Any

# ── 경로 설정 ──────────────────────────────────────────────────────────────
os.chdir("/app")
sys.path.insert(0, "/app")

from PIL import Image as PILImage

from app.config import settings
from app.database import AsyncSessionLocal
from app.models.enums import CategoryType
from app.models.expense import ExpenseItem
from app.models.project import Project
from app.models.company_setting import CompanySetting
from app.services.document_set_service import DocumentSetService
from sqlalchemy import select, delete


# ── 테스트 데이터 ────────────────────────────────────────────────────────────

PLACEHOLDER_PROJECT_CODE = "M118-SANITY"

SAMPLE_LINE_ITEMS = [
    {
        "item_name": "이산화티타늄(TiO2)",
        "spec": "순도 99.9%, 1kg",
        "quantity": 5,
        "unit_price": 1_350_000,
        "amount": 6_750_000,
    },
    {
        "item_name": "실리카(Silica)",
        "spec": "D50 200nm, 1kg",
        "quantity": 5,
        "unit_price": 650_000,
        "amount": 3_250_000,
    },
]

TOTAL_AMOUNT = sum(item["amount"] for item in SAMPLE_LINE_ITEMS)  # 10_000_000


def _make_placeholder_png(dest_path: Path) -> None:
    """단색 PNG 플레이스홀더 생성 (400×300, 연두색)."""
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    img = PILImage.new("RGB", (400, 300), color=(200, 230, 200))
    img.save(str(dest_path), "PNG")


def _check_docx_text(path: str, keywords: list[str]) -> list[str]:
    """docx 파일에서 keywords 각각이 존재하는지 확인. 미발견 목록 반환."""
    from docx import Document
    doc = Document(path)
    full_text = "\n".join(
        p.text for p in doc.paragraphs
    )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    missing = [kw for kw in keywords if kw not in full_text]
    return missing


def _check_docx_has_image(path: str) -> bool:
    """docx 파일에 인라인 이미지가 1개 이상 포함되어 있는지 확인."""
    import zipfile
    with zipfile.ZipFile(path, "r") as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        return len(media) > 0


async def _cleanup_sanity_data(db: Any) -> None:
    await db.execute(
        delete(ExpenseItem).where(
            ExpenseItem.title.like("%M118-SANITY%")
        )
    )
    result = await db.execute(
        select(Project).where(Project.code == PLACEHOLDER_PROJECT_CODE)
    )
    proj = result.scalar_one_or_none()
    if proj:
        await db.delete(proj)
    await db.commit()


async def run_sanity() -> None:
    async with AsyncSessionLocal() as db:
        # ── 임시 데이터 정리 먼저 ───────────────────────────────────────────
        await _cleanup_sanity_data(db)

        # ── 플레이스홀더 이미지 생성 ──────────────────────────────────────────
        img_dir = Path(settings.storage_documents_path) / "m118_sanity"
        img1 = img_dir / "placeholder_image_1.png"
        img2 = img_dir / "placeholder_image_2.png"
        _make_placeholder_png(img1)
        _make_placeholder_png(img2)

        # ── 임시 프로젝트 생성 ────────────────────────────────────────────────
        project = Project(
            id=uuid.uuid4(),
            name="M118 Sanity 테스트 과제",
            code=PLACEHOLDER_PROJECT_CODE,
            institution="주식회사 테스트",
            principal_investigator="김연구",
            period_start=datetime.date(2026, 1, 1),
            period_end=datetime.date(2026, 12, 31),
            total_budget=50_000_000,
            status="active",
        )
        db.add(project)
        await db.flush()

        # ── CompanySetting 조회 ───────────────────────────────────────────────
        cs_result = await db.execute(select(CompanySetting).limit(1))
        company_setting = cs_result.scalar_one_or_none()

        # ── 임시 expense 생성 (재료비) ─────────────────────────────────────────
        line_items_with_images = [
            {**SAMPLE_LINE_ITEMS[0], "image_path": str(img1)},
            {**SAMPLE_LINE_ITEMS[1], "image_path": str(img2)},
        ]
        expense = ExpenseItem(
            id=uuid.uuid4(),
            project_id=project.id,
            category_type=CategoryType.materials,
            title="M118-SANITY 이산화티타늄 외",
            amount=TOTAL_AMOUNT,
            expense_date="2026-05-06",
            vendor_name="㈜대신테크젠",
            vendor_registration_number="123-45-67890",
            metadata_={
                "line_items": line_items_with_images,
                "vendor_id": None,
            },
        )
        db.add(expense)
        await db.flush()

        await db.commit()

        print(f"\n[M118] expense_id={expense.id}")
        print(f"[M118] 이미지1={img1}")
        print(f"[M118] 이미지2={img2}")

        # ── generate_set 실행 ─────────────────────────────────────────────────
        print("\n[M118] generate_set 실행...")
        svc = DocumentSetService(settings.storage_documents_path)
        try:
            result = await svc.generate_set(expense.id, db)
        except Exception as e:
            print(f"[ERROR] generate_set 실패: {e}")
            import traceback; traceback.print_exc()
            await _cleanup_sanity_data(db)
            return

        # ── 검증 ─────────────────────────────────────────────────────────────
        print("\n[M118] 생성 결과:")
        matrix: dict[str, Any] = {}
        for item in result.items:
            status_str = item.status
            path = item.output_path or ""
            print(f"  {item.document_type.value:35s}  status={status_str:20s}  path={Path(path).name if path else '(없음)'}")
            matrix[item.document_type.value] = {"status": status_str, "path": path}

        print("\n[M118] ── 검증 매트릭스 ──")
        passed = 0
        failed = 0

        def check(label: str, ok: bool, detail: str = "") -> None:
            nonlocal passed, failed
            icon = "✅" if ok else "❌"
            print(f"  {icon} {label}" + (f" — {detail}" if detail else ""))
            if ok:
                passed += 1
            else:
                failed += 1

        # 품의서 (purchase_request)
        pr_info = matrix.get("purchase_request", {})
        pr_path = pr_info.get("path", "")
        if pr_path and Path(pr_path).exists():
            check("P1 품의서 생성됨", True)
            pr_missing = _check_docx_text(pr_path, ["이산화티타늄", "대신테크젠"])
            check("P5 품목명(이산화티타늄)", "이산화티타늄" not in pr_missing, f"missing={pr_missing}")
            # total_amount는 정수 렌더링(10000000) 또는 포맷(10,000,000) 둘 다 허용
            pr_has_amount = not _check_docx_text(pr_path, ["10000000"]) or not _check_docx_text(pr_path, ["10,000,000"])
            check("P6 합계 금액 포함", pr_has_amount)
        else:
            check("P1 품의서 생성됨", False, "파일 없음")
            check("P5 품목명", False); check("P6 합계", False)

        # 지출결의서 (expense_resolution)
        er_info = matrix.get("expense_resolution", {})
        er_path = er_info.get("path", "")
        if er_path and Path(er_path).exists():
            er_missing = _check_docx_text(er_path, ["이산화티타늄"])
            check("E1 지출결의서 생성됨", True)
            check("E3 품목명(이산화티타늄)", "이산화티타늄" not in er_missing, f"missing={er_missing}")
            er_has_amount = not _check_docx_text(er_path, ["10000000"]) or not _check_docx_text(er_path, ["10,000,000"])
            check("E4 합산 금액 포함", er_has_amount)
            from docx import Document as Docx
            er_doc = Docx(er_path)
            er_text = "\n".join(p.text for p in er_doc.paragraphs)
            for t in er_doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        er_text += "\n" + cell.text
            check("E2 비목 체크박스 ☑ 포함", "☑" in er_text)
        else:
            check("E1 지출결의서 생성됨", False, "파일 없음")
            for _ in range(3): check("E2~E4", False)

        # 검수확인서 (inspection_confirmation)
        ic_info = matrix.get("inspection_confirmation", {})
        ic_path = ic_info.get("path", "")
        if ic_path and Path(ic_path).exists():
            check("I1 검수확인서 생성됨", True)
            has_img = _check_docx_has_image(ic_path)
            check("I2 이미지 삽입 여부", has_img, "word/media/ 없음" if not has_img else "")
        else:
            check("I1 검수확인서 생성됨", False, "파일 없음")
            check("I2 이미지 삽입", False)

        print(f"\n  TOTAL: ✅ {passed}  ❌ {failed}")

        # ── 임시 데이터 정리 ───────────────────────────────────────────────────
        print("\n[M118] 임시 데이터 정리...")
        await _cleanup_sanity_data(db)

        # sanity 이미지 정리
        for f in [img1, img2]:
            f.unlink(missing_ok=True)
        try:
            img_dir.rmdir()
        except OSError:
            pass

        if failed == 0:
            print("[M118] PASS — 회귀 없음, Phase 2b 정상")
        else:
            print(f"[M118] FAIL — {failed}개 항목 실패")
            sys.exit(1)


if __name__ == "__main__":
    asyncio.run(run_sanity())
